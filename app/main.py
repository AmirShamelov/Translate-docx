from fastapi import FastAPI, UploadFile, File, Query
from fastapi.responses import StreamingResponse
from docx import Document
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from io import BytesIO
from urllib.parse import quote
import aiohttp

app = FastAPI()

# АПИ переводчик
API_BASE = "http://34.9.223.19:8000/translate"


# Отправляет текст в сторонний АПИ для перевода
async def translate_text(text: str, direction: str) -> str:
    # param text: Оригинальный текст
    # param direction: Выбор перевода (ru-kk, kk-ru)
    url = f"{API_BASE}/{direction}/"
    async with aiohttp.ClientSession() as session:
        async with session.post(url, json={"text": text}) as response:
            if response.status == 200:
                result = await response.json()
                return result.get("translated_text", "")  # Возвращает переведенный текст
            return text  # Если АПИ не ответил возвращаем оригинал текста


# Генератор, который возвращает поочередно блоки в документе (Параграфы и таблицы)
def iter_block_items(parent):
    for child in parent.element.body:
        if child.tag.endswith("}p"):
            yield Paragraph(child, parent)
        elif child.tag.endswith("}tbl"):
            yield Table(child, parent)


# Переводит текст внутри параграфа и заменяет его новым переведенным текстом
async def translate_paragraph(paragraph: Paragraph, direction: str):
    text = paragraph.text
    if not text.strip():
        return  # пропуск пустых графов

    translated = await translate_text(text, direction)

    # Удаляем все части текста с разным форматированием (run)
    for _ in range(len(paragraph.runs)):
        paragraph.runs[0].clear()

    # Добавляем один новый run с переведенным текстом
    run = paragraph.add_run(translated)

    # Копируем базовое форматирование из первого run
    if paragraph.runs:
        style_source = paragraph.runs[0]
        run.bold = style_source.bold
        run.italic = style_source.italic
        run.underline = style_source.underline


# Переводит весь текст внутри таблицы
async def translate_table(table: Table, direction: str):
    for row in table.rows:
        for cell in row.cells:
            await translate_cell(cell, direction)


# Переводит все параграфы и таблицы внутри одной ячейки таблицы
async def translate_cell(cell: _Cell, direction: str):
    # Переводим все параграфы внутри ячейки
    for paragraph in cell.paragraphs:
        await translate_paragraph(paragraph, direction)

    # Переводим все таблицы внутри ячейки
    for table in cell.tables:
        await translate_table(table, direction)


# Основной эндпоинт. Принимает док файл, переводит его содержимое и возвращает новый док файл с сохранением структуры
@app.post("/translate-docx/")
async def translate_docx(
        file: UploadFile = File(...),
        direction: str = Query(..., regex="^(ru-kk|kk-ru)$", description="Направление перевода: ru-kk или kk-ru")
):
    contents = await file.read()
    doc = Document(BytesIO(contents))

    # Каждый уровень блока в документе
    for item in iter_block_items(doc):
        if isinstance(item, Paragraph):
            await translate_paragraph(item, direction)
        elif isinstance(item, Table):
            await translate_table(item, direction)

    # Сохраняем переведённый документ в память
    output = BytesIO()
    doc.save(output)
    output.seek(0)

    # Корректное имя файла для загрузки (использовал библиотеку quote чтобы правильно форматировался название файла)
    filename_utf8 = f"translated_{file.filename}"
    disposition = f"attachment; filename*=UTF-8''{quote(filename_utf8)}"

    return StreamingResponse(
        output,
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
        headers={"Content-Disposition": disposition},
    )
